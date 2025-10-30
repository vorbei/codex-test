#!/usr/bin/env python3
"""
PDF转DOCX工具 (改进版)
使用PyMuPDF提取内容，python-docx生成文档
"""

import argparse
import os
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError as e:
    print(f"错误: 缺少必要的库。请运行: pip install PyMuPDF python-docx")
    print(f"详细错误: {e}")
    sys.exit(1)


def extract_text_from_pdf(pdf_path, start_page=0, end_page=None):
    """
    从PDF中提取文本内容

    参数:
        pdf_path: PDF文件路径
        start_page: 起始页码（从0开始）
        end_page: 结束页码（None表示到最后一页）

    返回:
        list: 每页的文本内容列表
    """
    try:
        doc = fitz.open(pdf_path)
        pages_text = []

        # 确定页面范围
        if end_page is None:
            end_page = len(doc)
        else:
            end_page = min(end_page, len(doc))

        # 提取每页文本
        for page_num in range(start_page, end_page):
            page = doc[page_num]
            text = page.get_text()
            pages_text.append({
                'page_num': page_num + 1,
                'text': text
            })

        doc.close()
        return pages_text

    except Exception as e:
        print(f"提取PDF文本失败: {str(e)}")
        return None


def create_docx_from_text(text_data, output_path):
    """
    从文本数据创建DOCX文档

    参数:
        text_data: 页面文本数据列表
        output_path: 输出DOCX文件路径
    """
    try:
        doc = Document()

        for page_data in text_data:
            # 添加页码标题
            if len(text_data) > 1:  # 只在多页时显示页码
                heading = doc.add_heading(f'Page {page_data["page_num"]}', level=2)
                heading.runs[0].font.size = Pt(14)

            # 添加页面文本
            text = page_data['text'].strip()
            if text:
                # 按段落分割
                paragraphs = text.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        p.style = 'Normal'
                        # 设置字体
                        for run in p.runs:
                            run.font.size = Pt(11)

            # 在页面之间添加分页符（除了最后一页）
            if page_data != text_data[-1] and len(text_data) > 1:
                doc.add_page_break()

        # 保存文档
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"创建DOCX文档失败: {str(e)}")
        return False


def convert_pdf_to_docx(pdf_path, docx_path=None, start_page=0, end_page=None):
    """
    将PDF文件转换为DOCX格式

    参数:
        pdf_path: PDF文件路径
        docx_path: 输出的DOCX文件路径（可选，默认与PDF同名）
        start_page: 起始页码（从0开始）
        end_page: 结束页码（None表示到最后一页）

    返回:
        bool: 转换是否成功
    """
    # 检查输入文件是否存在
    if not os.path.exists(pdf_path):
        print(f"错误: PDF文件不存在: {pdf_path}")
        return False

    # 如果未指定输出路径，使用相同的文件名，只改变扩展名
    if docx_path is None:
        docx_path = str(Path(pdf_path).with_suffix('.docx'))

    try:
        print(f"开始转换: {pdf_path} -> {docx_path}")

        # 提取PDF文本
        print("  [1/2] 提取PDF文本...")
        text_data = extract_text_from_pdf(pdf_path, start_page, end_page)

        if text_data is None:
            return False

        if not text_data:
            print("警告: PDF中没有提取到文本内容")
            return False

        print(f"  成功提取 {len(text_data)} 页内容")

        # 创建DOCX文档
        print("  [2/2] 生成DOCX文档...")
        if create_docx_from_text(text_data, docx_path):
            print(f"转换成功! 输出文件: {docx_path}")
            return True
        else:
            return False

    except Exception as e:
        print(f"转换失败: {str(e)}")
        return False


def batch_convert(input_dir, output_dir=None):
    """
    批量转换目录中的所有PDF文件

    参数:
        input_dir: 包含PDF文件的目录
        output_dir: 输出目录（可选，默认与输入目录相同）
    """
    if not os.path.isdir(input_dir):
        print(f"错误: 目录不存在: {input_dir}")
        return

    # 如果未指定输出目录，使用输入目录
    if output_dir is None:
        output_dir = input_dir
    else:
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

    # 查找所有PDF文件
    pdf_files = list(Path(input_dir).glob("*.pdf"))

    if not pdf_files:
        print(f"在目录 {input_dir} 中未找到PDF文件")
        return

    print(f"找到 {len(pdf_files)} 个PDF文件\n")

    success_count = 0
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"[{i}/{len(pdf_files)}] 转换 {pdf_file.name}")
        output_path = os.path.join(output_dir, pdf_file.stem + '.docx')
        if convert_pdf_to_docx(str(pdf_file), output_path):
            success_count += 1
        print()

    print(f"批量转换完成: {success_count}/{len(pdf_files)} 个文件成功转换")


def main():
    """命令行接口"""
    parser = argparse.ArgumentParser(
        description='将PDF文件转换为DOCX格式（改进版 - 使用文本提取）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s input.pdf                    # 转换单个文件
  %(prog)s input.pdf -o output.docx     # 指定输出文件名
  %(prog)s input.pdf -s 0 -e 5          # 只转换第1到第5页
  %(prog)s --batch ./pdfs               # 批量转换目录中的所有PDF
  %(prog)s --batch ./pdfs -o ./docx     # 批量转换并指定输出目录

注意:
  此版本提取PDF中的文本内容并重新格式化为DOCX。
  复杂的格式、图片和表格可能无法完美保留。
        """
    )

    parser.add_argument('input',
                       help='输入的PDF文件路径或目录（批量模式）')
    parser.add_argument('-o', '--output',
                       help='输出的DOCX文件路径或目录（批量模式）')
    parser.add_argument('-s', '--start-page',
                       type=int,
                       default=0,
                       help='起始页码（从0开始，默认为0）')
    parser.add_argument('-e', '--end-page',
                       type=int,
                       help='结束页码（默认为最后一页）')
    parser.add_argument('-b', '--batch',
                       action='store_true',
                       help='批量转换模式：转换目录中的所有PDF文件')

    args = parser.parse_args()

    # 批量转换模式
    if args.batch:
        batch_convert(args.input, args.output)
    else:
        # 单文件转换模式
        result = convert_pdf_to_docx(
            args.input,
            args.output,
            args.start_page,
            args.end_page
        )
        sys.exit(0 if result else 1)


if __name__ == "__main__":
    main()
